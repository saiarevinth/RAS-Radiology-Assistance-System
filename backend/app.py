from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import torch
import numpy as np
import cv2
import io
import base64
from PIL import Image
from model import ResUNet50   # make sure model.py defines this0
import json
from datetime import datetime
import os
from pdf_extractor import extract_medical_fields_from_pdf
from werkzeug.utils import secure_filename

# Device configuration

DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# Initialize Flask app
app = Flask(__name__)

# Load configuration
from config import Config
app.config.from_object(Config)

# Enable credentialed CORS so we can use HttpOnly cookies from the frontend
CORS(
    app,
    supports_credentials=True,
    resources={r"/*": {"origins": ["http://localhost:3000", "http://127.0.0.1:3000"]}},
)

# Initialize database
from models import db
db.init_app(app)

# Register blueprints
from auth import auth_bp, require_auth  # noqa: E402
from database import database_bp  # noqa: E402

app.register_blueprint(auth_bp)
app.register_blueprint(database_bp)

# --- Patient Intake API (add assignedDoctorId support) ---
from models import db, Patient, PatientIntake, User, MedicalReport


# --- Doctor profile patient lists ---
@app.route('/api/doctor/<int:doctor_id>/assigned-patients', methods=['GET'])
def get_assigned_patients(doctor_id):
    # Get all intakes assigned to this doctor
    intakes = PatientIntake.query.filter_by(assigned_doctor_id=doctor_id).order_by(PatientIntake.created_at.desc()).all()
    # Map patient_id to latest intake (by created_at)
    latest_intake_by_patient = {}
    for intake in intakes:
        pid = intake.patient_id
        if pid not in latest_intake_by_patient or intake.created_at > latest_intake_by_patient[pid].created_at:
            latest_intake_by_patient[pid] = intake
    patient_ids = list(latest_intake_by_patient.keys())
    patients = Patient.query.filter(Patient.id.in_(patient_ids)).all() if patient_ids else []
    # Combine patient and latest intake info
    assigned_patients = []
    for patient in patients:
        intake = latest_intake_by_patient.get(patient.id)
        patient_dict = patient.to_dict()
        if intake:
            intake_dict = intake.to_dict()
            # Merge relevant intake fields (including high_priority) into patient_dict
            patient_dict['intake'] = intake_dict
            patient_dict['high_priority'] = intake_dict.get('high_priority', False)
        else:
            patient_dict['intake'] = None
            patient_dict['high_priority'] = False
        assigned_patients.append(patient_dict)
    return jsonify({'assigned_patients': assigned_patients})


@app.route('/api/doctor/<int:doctor_id>/attended-patients', methods=['GET'])
def get_attended_patients(doctor_id):
    # Get all patients attended by this doctor via MedicalReport
    report_patient_ids = set(report.patient_id for report in MedicalReport.query.filter_by(doctor_id=doctor_id).all())
    # Get all patients manually marked as attended (new logic)
    manual_attended = PatientIntake.query.filter_by(assigned_doctor_id=doctor_id).filter(PatientIntake.report_content == 'attended').all()
    manual_patient_ids = set(intake.patient_id for intake in manual_attended)
    all_patient_ids = list(report_patient_ids.union(manual_patient_ids))
    patients = Patient.query.filter(Patient.id.in_(all_patient_ids)).all() if all_patient_ids else []
    return jsonify({'attended_patients': [p.to_dict() for p in patients]})

@app.route('/api/doctor/<int:doctor_id>/attend-patient', methods=['POST'])
def attend_patient(doctor_id):
    """Mark an assigned patient as attended by this doctor (manual or via extraction)"""
    user = require_auth(request)
    if not user or user.role != 'doctor' or user.id != doctor_id:
        return jsonify({'error': 'Unauthorized'}), 401
    data = request.get_json()
    patient_id = data.get('patient_id')
    if not patient_id:
        return jsonify({'error': 'Missing patient_id'}), 400
    try:
        # Find the latest intake for this patient assigned to this doctor
        intake = PatientIntake.query.filter_by(patient_id=patient_id, assigned_doctor_id=doctor_id).order_by(PatientIntake.created_at.desc()).first()
        if not intake:
            return jsonify({'error': 'No such assigned patient'}), 404
        # Mark as attended (use report_content field as a flag)
        intake.report_content = 'attended'
        intake.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'message': 'Patient marked as attended'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to mark as attended: {str(e)}'}), 500

@app.route('/api/doctor/<int:doctor_id>/undo-attend-patient', methods=['POST'])
def undo_attend_patient(doctor_id):
    """Undo attending a patient, relisting as assigned only."""
    user = require_auth(request)
    if not user or user.role != 'doctor' or user.id != doctor_id:
        return jsonify({'error': 'Unauthorized'}), 401
    data = request.get_json()
    patient_id = data.get('patient_id')
    if not patient_id:
        return jsonify({'error': 'Missing patient_id'}), 400
    try:
        intake = PatientIntake.query.filter_by(patient_id=patient_id, assigned_doctor_id=doctor_id).order_by(PatientIntake.created_at.desc()).first()
        if not intake:
            return jsonify({'error': 'No such assigned patient'}), 404
        # Undo attended: clear the attended flag if it was set
        if intake.report_content == 'attended':
            intake.report_content = None
            intake.updated_at = datetime.utcnow()
            db.session.commit()
        return jsonify({'success': True, 'message': 'Patient relisted as assigned'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to undo attend: {str(e)}'}), 500


# Simple model cache so you don't reload on every request
_model_cache = {}

def get_model(model_path):
    if model_path in _model_cache:
        return _model_cache[model_path]

    model = ResUNet50(out_channels=1, pretrained=False)
    state = torch.load(model_path, map_location=DEVICE)

    if isinstance(state, dict) and "state_dict" in state:
        model.load_state_dict(state["state_dict"])
    else:
        try:
            model.load_state_dict(state)
        except Exception:
            model = state

    model.to(DEVICE)
    model.eval()
    _model_cache[model_path] = model
    return model

def preprocess_image(image_data, target_size=(256, 256)):
    image = Image.open(io.BytesIO(image_data))
    if hasattr(image, "n_frames") and image.n_frames > 1:
        image.seek(0)
    image = np.array(image)

    if len(image.shape) == 2:
        image = cv2.cvtColor(image, cv2.COLOR_GRAY2RGB)
    elif image.shape[2] > 3:
        image = image[:, :, :3]

    image_resized = cv2.resize(image, target_size)
    image_normalized = image_resized.astype(np.float32) / 255.0
    image_tensor = torch.tensor(image_normalized).permute(2, 0, 1).unsqueeze(0)
    return image, image_tensor

def predict(model, image_tensor):
    with torch.no_grad():
        image_tensor = image_tensor.to(DEVICE)
        out = model(image_tensor)
        out = torch.sigmoid(out)
        out = (out > 0.5).float()
    return out

def overlay_mask(original_image, mask, alpha=0.5):
    if len(original_image.shape) == 2:
        original_image = cv2.cvtColor(original_image, cv2.COLOR_GRAY2RGB)
    if original_image.max() > 1.0:
        original_image = original_image.astype(np.float32) / 255.0

    mask_resized = cv2.resize(mask, (original_image.shape[1], original_image.shape[0]), interpolation=cv2.INTER_NEAREST)
    colored_mask = np.zeros_like(original_image, dtype=np.float32)
    colored_mask[mask_resized > 0] = [0, 0, 1.0]  # Blue mask
    result = cv2.addWeighted(original_image, 1.0, colored_mask, alpha, 0)
    return (result * 255).astype(np.uint8)

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for frontend to verify backend connectivity"""
    try:
        # Check database connection
        db.session.execute('SELECT 1')
        return jsonify({
            "status": "healthy",
            "database": "connected",
            "timestamp": datetime.utcnow().isoformat()
        }), 200
    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "database": "disconnected",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        }), 500

@app.route('/upload-pdf', methods=['POST'])
def upload_pdf():
    """Upload PDF file for patient intake"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Check if it's a PDF
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({"error": "Only PDF files are allowed"}), 400
        
        # Create uploads directory if it doesn't exist
        upload_folder = 'uploads'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        
        # Generate unique filename
        import uuid
        unique_filename = f"{uuid.uuid4()}_{secure_filename(file.filename)}"
        file_path = os.path.join(upload_folder, unique_filename)
        
        # Save file
        file.save(file_path)
        
        return jsonify({
            "success": True,
            "message": "PDF uploaded successfully",
            "filename": unique_filename,
            "file_path": file_path
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Failed to upload PDF: {str(e)}"}), 500

@app.route('/uploads/<filename>')
def serve_uploaded_file(filename):
    """Serve uploaded files"""
    try:
        return send_from_directory('uploads', filename)
    except Exception as e:
        return jsonify({"error": f"File not found: {str(e)}"}), 404

def png_bytes_to_datauri(png_bytes):
    b64 = base64.b64encode(png_bytes).decode("ascii")
    return f"data:image/png;base64,{b64}"



# PDF extraction endpoint for patient intake
@app.route("/extract-pdf", methods=["POST"])
def extract_pdf():
    """
    Extract medical information from uploaded PDF files.
    Expects multipart form data with 'file' field containing PDF.
    """
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        # Check if file is empty
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Check file type
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({"error": "Only PDF files are supported"}), 400
        
        # Read file bytes
        file_bytes = file.read()
        
        # Check file size (limit to 10MB)
        if len(file_bytes) > 10 * 1024 * 1024:
            return jsonify({"error": "File size too large. Maximum 10MB allowed."}), 400
        
        # Extract medical fields using the PDF extractor
        extracted_data = extract_medical_fields_from_pdf(file_bytes)
        
        # Return extracted data
        return jsonify({
            "success": True,
            "data": extracted_data,
            "message": "PDF processed successfully"
        }), 200
        
    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Failed to process PDF: {str(e)}"
        }), 500

# API endpoint for segmentation
@app.route("/segment", methods=["POST"])
def segment():
    # Require authentication
    user = require_auth(request)
    if not user:
        return jsonify({"error": "unauthorized"}), 401
    if "image" not in request.files:
        return jsonify({"error": "no image file"}), 400

    model_path = request.form.get("model_path", "resunet50_brain_segmentation.pth")
    image_file = request.files["image"]

    try:
        model = get_model(model_path)
    except Exception as e:
        return jsonify({"error": f"failed to load model: {e}"}), 500

    image_bytes = image_file.read()
    original_image, image_tensor = preprocess_image(image_bytes)
    mask_pred = predict(model, image_tensor)
    mask_numpy = mask_pred[0, 0].cpu().numpy()

    overlay_img = overlay_mask(original_image, mask_numpy)
    tumor_pixels = int(np.sum(mask_numpy > 0.5))
    total_pixels = int(mask_numpy.size)
    affected_percentage = round((tumor_pixels / total_pixels) * 100, 2)

    success, png = cv2.imencode(".png", overlay_img)
    if not success:
        return jsonify({"error": "failed to encode image"}), 500

    data_uri = png_bytes_to_datauri(png.tobytes())

    return jsonify({
        "image_data_uri": data_uri,
        "affected_percentage": affected_percentage
    })

@app.route("/export-report", methods=["POST"])
def export_report():
    # Require authentication
    user = require_auth(request)
    if not user:
        return jsonify({"error": "unauthorized"}), 401
    
    try:
        data = request.get_json()
        format_type = data.get("format", "pdf")
        report_data = data.get("reportData", {})
        
        if format_type == "html":
            return generate_html_report(report_data)
        elif format_type == "pdf":
            return generate_pdf_report(report_data)
        elif format_type == "docx":
            return generate_docx_report(report_data)
        else:
            return jsonify({"error": "Unsupported format"}), 400
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to export report: {str(e)}"}), 500

def generate_pdf_report(report_data):
    """Generate PDF report (simplified - returns HTML for now)"""
    # For now, return HTML that can be converted to PDF by the browser
    # In production, you'd use a library like reportlab or weasyprint
    html_content = generate_html_content(report_data)
    
    response = app.response_class(
        response=html_content,
        status=200,
        mimetype='text/html'
    )
    filename = generate_safe_filename(report_data.get('patientName', 'patient'), 'pdf')
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def generate_docx_report(report_data):
    """Generate DOCX report using python-docx; fallback to HTML-as-DOC if unavailable"""
    # Try importing python-docx. If it fails, fallback immediately.
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        import traceback
        traceback.print_exc()
        html_content = generate_html_content(report_data)
        response = app.response_class(
            response=html_content,
            status=200,
            mimetype='application/msword'
        )
        filename = generate_safe_filename(report_data.get('patientName', 'patient'), 'doc')
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    import re

    # Build the document
    doc = Document()

    def add_heading(doc, text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 14)
        return p

    def add_kv(paragraph_text):
        p = doc.add_paragraph()
        if ":" in paragraph_text:
            key, value = paragraph_text.split(":", 1)
            run_b = p.add_run(f"{key}:")
            run_b.bold = True
            p.add_run(value)
        else:
            p.add_run(paragraph_text)

    def add_rich_paragraph(text: str):
        # Handle bullets like lines starting with '* '
        lines = (text or '').split('\n')
        for line in lines:
            if re.match(r"^\s*\*\s+", line):
                p = doc.add_paragraph(style='List Bullet')
                content = re.sub(r"^\s*\*\s+", "", line)
            else:
                p = doc.add_paragraph()
                content = line
            # Convert **bold** segments into bold runs
            parts = re.split(r"(\*\*.+?\*\*)", content)
            for part in parts:
                if not part:
                    continue
                m = re.match(r"^\*\*(.+)\*\*$", part)
                if m:
                    run = p.add_run(m.group(1))
                    run.bold = True
                else:
                    p.add_run(part)

    # Header
    title = doc.add_paragraph()
    tr = title.add_run("Medical Imaging Report")
    tr.bold = True
    tr.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

    # Patient Information
    add_heading(doc, "Patient Information", level=1)
    add_kv(f"Patient Name: {report_data.get('patientName', 'N/A')}")
    add_kv(f"Referring Physician: {report_data.get('patientInfo', {}).get('referringPhysician', 'N/A')}")
    add_kv(f"Report Generated By: {report_data.get('doctorName', 'N/A')} ({report_data.get('doctorSpecialty', 'N/A')})")
    add_kv(f"Chief Complaint: {report_data.get('patientInfo', {}).get('chiefComplaint', 'N/A')}")
    add_kv(f"Affected Area: {report_data.get('affectedPercentage', 'N/A')}%")
    add_kv(f"Report Status: {'Edited' if report_data.get('isEdited') else 'AI Generated'}")

    # Clinical History
    add_heading(doc, "Clinical History", level=1)
    add_kv(f"Medical History: {report_data.get('patientInfo', {}).get('medicalHistory', 'N/A')}")
    add_kv(f"Current Medications: {report_data.get('patientInfo', {}).get('currentMedications', 'N/A')}")
    add_kv(f"Known Allergies: {report_data.get('patientInfo', {}).get('knownAllergies', 'N/A')}")
    add_kv(f"Family History: {report_data.get('patientInfo', {}).get('familyHistory', 'N/A')}")

    # AI Analysis Report
    add_heading(doc, "AI Analysis Report", level=1)
    content = report_data.get('content', 'No report content available') or ''
    # Split into paragraphs and structure bullets
    for para in re.split(r"\n\s*\n", content.strip()):
        lines = para.strip().split('\n')
        for line in lines:
            if line.strip().startswith('* '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                # Heading-like bold
                p = doc.add_paragraph()
                run = p.add_run(line.strip().strip('*'))
                run.bold = True
            else:
                doc.add_paragraph(line.strip())

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("This report was generated using AI-assisted medical imaging analysis.")
    doc.add_paragraph("Please review all findings with qualified medical professionals.")

    # Stream as response
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    response = app.response_class(
        response=mem.getvalue(),
        status=200,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = generate_safe_filename(report_data.get('patientName', 'patient'), 'docx')
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def generate_html_content(report_data):
    """Generate HTML content for reports"""
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Medical Report - {report_data.get('patientName', 'Unknown')}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            .header {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 30px; }}
            .section {{ margin-bottom: 25px; }}
            .section h2 {{ color: #2c5aa0; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
            .info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
            .patient-info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; }}
            .report-content {{ white-space: pre-wrap; line-height: 1.6; }}
            .footer {{ margin-top: 40px; text-align: center; font-size: 12px; color: #666; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>Medical Imaging Report</h1>
            <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
        </div>
        
        <div class="section">
            <h2>Patient Information</h2>
            <div class="info-grid">
                <div class="patient-info">
                    <strong>Patient Name:</strong> {report_data.get('patientName', 'N/A')}<br>
                    <strong>Referring Physician:</strong> {report_data.get('patientInfo', {}).get('referringPhysician', 'N/A')}<br>
                    <strong>Report Generated By:</strong> {report_data.get('doctorName', 'N/A')} ({report_data.get('doctorSpecialty', 'N/A')})
                </div>
                <div class="patient-info">
                    <strong>Chief Complaint:</strong> {report_data.get('patientInfo', {}).get('chiefComplaint', 'N/A')}<br>
                    <strong>Affected Area:</strong> {report_data.get('affectedPercentage', 'N/A')}%<br>
                    <strong>Report Status:</strong> {'Edited' if report_data.get('isEdited') else 'AI Generated'}
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>Clinical History</h2>
            <div class="patient-info">
                <strong>Medical History:</strong> {report_data.get('patientInfo', {}).get('medicalHistory', 'N/A')}<br><br>
                <strong>Current Medications:</strong> {report_data.get('patientInfo', {}).get('currentMedications', 'N/A')}<br><br>
                <strong>Known Allergies:</strong> {report_data.get('patientInfo', {}).get('knownAllergies', 'N/A')}<br><br>
                <strong>Family History:</strong> {report_data.get('patientInfo', {}).get('familyHistory', 'N/A')}
            </div>
        </div>
        
        <div class="section">
            <h2>AI Analysis Report</h2>
            <div class="report-content">{report_data.get('content', 'No report content available')}</div>
        </div>
        
        <div class="footer">
            <p>This report was generated using AI-assisted medical imaging analysis.</p>
            <p>Please review all findings with qualified medical professionals.</p>
        </div>
    </body>
    </html>
    """

def generate_html_report(report_data):
    """Generate HTML report"""
    html_content = generate_html_content(report_data)
    
    response = app.response_class(
        response=html_content,
        status=200,
        mimetype='text/html'
    )
    filename = generate_safe_filename(report_data.get('patientName', 'patient'), 'html')
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def generate_safe_filename(name, ext):
    import re
    safe = re.sub(r'[^A-Za-z0-9._-]', '_', str(name or 'patient'))
    return f"{safe}_report.{ext}"

if __name__ == "__main__":
    from llama_warmup import warmup_llama_model
    warmup_llama_model()
    app.run(host="0.0.0.0", port=5000, debug=True)
